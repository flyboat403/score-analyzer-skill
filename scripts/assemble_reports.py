#!/usr/bin/env python3
"""
Standalone report assembly script.
Creates Word and HTML reports with embedded charts.
"""

import pandas as pd
import os
import sys
import argparse
import base64
import re
import zipfile
from typing import Dict, Optional
from docx import Document
from docx.shared import Inches


def dataframe_to_markdown(df, index=False):
    """Convert DataFrame to Markdown table format."""
    df = df.copy()

    if index:
        df = df.reset_index()

    headers = df.columns.tolist()
    rows = [[str(val) for val in row.values] for _, row in df.iterrows()]

    md_lines = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join(["---"] * len(headers)) + "|"
    ]
    for row in rows:
        md_lines.append("| " + " | ".join(row) + " |")

    return "\n".join(md_lines)


def parse_markdown_simple(markdown_text):
    """Simple markdown parser for section extraction."""
    sections = []
    current_section = {'title': '', 'content': ''}

    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            if current_section['content'] or current_section['title']:
                sections.append(current_section.copy())
                current_section = {'title': '', 'content': ''}
        elif line.startswith('### '):
            if current_section['content']:
                sections.append(current_section.copy())
            current_section = {'title': line.replace('### ', ''), 'content': ''}
        elif line.startswith('## '):
            if current_section['content']:
                sections.append(current_section.copy())
            current_section = {'title': line.replace('## ', ''), 'content': ''}
        elif line.startswith('# '):
            if current_section['content']:
                sections.append(current_section.copy())
            current_section = {'title': line.replace('# ', ''), 'content': ''}
        else:
            if not line.startswith('[PLOT:'):
                current_section['content'] += line + '\n'

    if current_section['content'] or current_section['title']:
        sections.append(current_section)

    return sections


def simple_markdown_to_html_with_tables(text):
    """Enhanced Markdown to HTML with table support."""
    lines = text.split('\n')
    html = []
    in_table = False
    table_rows = []
    in_code_block = False

    for line in lines:
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                html.append('<pre><code>')
            else:
                html.append('</code></pre>')
            continue

        if in_code_block:
            html.append(line + '\n')
            continue

        line_stripped = line.strip()
        is_table_row = line_stripped.startswith('|') and line_stripped.endswith('|')

        if is_table_row:
            if not in_table:
                in_table = True
                table_rows = []
                table_alignments = [] # Initialize alignment map
            
            # Parse cells
            raw_cells = [cell.strip() for cell in line_stripped.split('|')][1:-1]
            
            # Detect if this is the separator row (e.g., |:---|:---:| or |---|---|)
            # Regex to match only dashes and optional colons at start/end
            is_separator = all(re.match(r'^:?-+:?$', c) for c in raw_cells)
            
            if is_separator:
                # Extract alignment info
                align_map = []
                for cell in raw_cells:
                    if cell.startswith(':') and cell.endswith(':'):
                        align_map.append('text-align:center;')
                    elif cell.startswith(':'):
                        align_map.append('text-align:left;')
                    elif cell.endswith(':'):
                        align_map.append('text-align:right;')
                    else:
                        align_map.append('')
                table_alignments = align_map
            else:
                table_rows.append(raw_cells)
            continue

        # End of table block processing
        if in_table and table_rows:
            html.append('<table class="data-table"><thead><tr>')
            
            # Process Header (first row)
            if table_alignments and len(table_alignments) >= len(table_rows[0]):
                header_styles = table_alignments
            else:
                header_styles = [''] * len(table_rows[0])
                
            for idx, cell in enumerate(table_rows[0]):
                style = header_styles[idx] if idx < len(header_styles) else ''
                html.append(f'<th style="{style}">{cell.strip()}</th>')
            
            html.append('</tr></thead><tbody>')

            # Process Data rows
            for i in range(1, len(table_rows)):
                cells = table_rows[i]
                html.append('<tr>')
                for idx, cell in enumerate(cells):
                    style = header_styles[idx] if idx < len(header_styles) else ''
                    html.append(f'<td style="{style}">{cell.strip()}</td>')
                html.append('</tr>')
            
            html.append('</tbody></table>')
            in_table = False
            table_rows = []
            table_alignments = [] # Reset alignments

        if not line_stripped:
            html.append('<br>')
            continue

        if line_stripped.startswith('# '):
            html.append(f'<h1>{line_stripped[2:]}</h1>')
        elif line_stripped.startswith('## '):
            html.append(f'<h2>{line_stripped[3:]}</h2>')
        elif line_stripped.startswith('### '):
            html.append(f'<h3>{line_stripped[4:]}</h3>')
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            list_item = line_stripped[2:]
            list_item = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', list_item)
            list_item = re.sub(r'\*(.+?)\*', r'<em>\1</em>', list_item)
            list_item = re.sub(r'`(.+?)`', r'<code>\1</code>', list_item)
            html.append(f'<ul><li>{list_item}</li></ul>')
        elif line_stripped.startswith('> '):
            html.append(f'<blockquote>{line_stripped[2:]}</blockquote>')
        elif line_stripped.startswith('![') and 'PLOT:' in line_stripped:
            # Preserve image placeholders for later replacement
            html.append(line_stripped)
        elif line_stripped.startswith('[PLOT:'):
            # Preserve compact placeholders for later replacement
            html.append(line_stripped)
        else:
            paragraph = line_stripped
            paragraph = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', paragraph)
            paragraph = re.sub(r'\*(.+?)\*', r'<em>\1</em>', paragraph)
            paragraph = re.sub(r'`(.+?)`', r'<code>\1</code>', paragraph)
            if paragraph:
                html.append(f'<p>{paragraph}</p>')

    return ''.join(html)


def create_word_report(markdown_text, image_paths, output_path):
    """Create Word report from markdown and charts."""
    doc = Document()
    doc.add_heading('智能成绩分析报告', 0)

    sections = parse_markdown_simple(markdown_text)

    for section in sections:
        if section['title']:
            doc.add_heading(section['title'], level=1)

        content_lines = section['content'].split('\n')
        i = 0
        while i < len(content_lines):
            line = content_lines[i].strip()

            is_image_line = line.startswith('![') or line.startswith('[PLOT:')
            if is_image_line:
                chart_placeholder = line
                chart_name = ""

                if line.startswith('!['):
                    markdown_image = line
                    chart_name = markdown_image.split('(')[1].rstrip(')') if '(' in markdown_image else ''
                    chart_name_key = chart_name.replace('PLOT:', '').strip()
                else:
                    chart_name_key = line.replace('[PLOT:', '').replace(']', '').strip()

                chart_name_lower = chart_name_key.lower().replace(' ', '').replace('-', '')

                found = False
                for name, path in image_paths.items():
                    name_normalized = name.lower().replace(' ', '').replace('-', '')
                    if name_normalized == chart_name_lower:
                        if path and os.path.exists(path):
                            try:
                                doc.add_picture(path, width=Inches(6))
                                found = True
                            except Exception as e:
                                print(f"Warning: Failed to insert image: {e}")
                            break

                if not found:
                    print(f"Warning: Chart not found for placeholder: {chart_placeholder}")
            elif line:
                p = doc.add_paragraph(line)
                if line.startswith('**') and line.endswith('**'):
                    p.runs[0].bold = True
            i += 1

    doc.save(output_path)
    return output_path


def create_html_report(markdown_text, image_paths, output_path):
    """Create HTML report with embedded base64 charts."""
    html = '''
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  body { font-family: Arial, sans-serif, "Microsoft YaHei", "SimHei", sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; line-height: 1.8; color: #1e293b; }
  h1, h2, h3, h4, h5, h6 { color: #1e293b; margin-top: 1.8em; margin-bottom: 0.8em; font-weight: 600; }
  h1 { font-size: 2em; border-bottom: 2px solid #3b82f6; padding-bottom: 0.5em; }
  h2 { font-size: 1.6em; border-bottom: 1px solid #cbd5e1; padding-bottom: 0.3em; }
  h3 { font-size: 1.4em; }
  p { margin: 1em 0; }
  ul, ol { margin: 1em 0; padding-left: 1.8em; }
  li { margin: 0.5em 0; }
  strong { color: #2563eb; font-weight: 600; }
  em { color: #64748b; font-style: italic; }
  code { background: #f1f5f9; padding: 0.2em 0.4em; border-radius: 4px; font-family: monospace; font-size: 0.9em; }
  pre { background: #1e293b; color: #f1f5f9; padding: 1em; border-radius: 8px; overflow-x: auto; }
  pre code { background: none; padding: 0; color: inherit; }
  blockquote { border-left: 4px solid #3b82f6; padding-left: 1em; margin: 1em 0; color: #64748b; font-style: italic; }
  table { width: 100%; border-collapse: collapse; margin: 1.5em 0; }
  th, td { border: 1px solid #e2e8f0; padding: 12px; text-align: left; }
  th { background-color: #f1f5f9; font-weight: 600; }
  tr:nth-child(even) { background-color: #f8fafc; }
  img { max-width: 100%; height: auto; display: block; margin: 20px auto; border: 1px solid #e2e8f0; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }
</style>
</head>
<body>
'''

    content = simple_markdown_to_html_with_tables(markdown_text)

    alt_texts = {
        'DISTRIBUTION': '分数分布图',
        'COMPARISON': '分组对比图',
        'RADAR': '雷达图',
        'TREND': '成绩趋势图',
        'NORMAL': '正态分布图',
        'BOXPLOT': '分组稳定性图',
        'HEATMAP': '科目相关性图',
        'DEVIATION': '学生偏离度图',
        'TOP_BOTTOM': '尖子生对比图',
        'SCATTER': '总分关联散点图',
        'CDF': '分数累积分布图',
        'BOXPLOT_SUBJ': '科目箱线对比图',
        '散点图': '总分关联散点图',
        '累积分布': '分数累积分布图',
        '科目箱线图': '科目箱线对比图',
        '分数分布图': '分数分布图',
        '分组对比图': '分组对比图',
        '雷达图': '雷达图',
        '成绩趋势图': '成绩趋势图',
        '正态分布图': '正态分布图',
        '分组稳定性图': '分组稳定性图',
        '科目相关性图': '科目相关性图',
        '学生偏离度图': '学生偏离度图',
        '尖子生对比图': '尖子生对比图',
    }

    unique_files = {}
    for name, path in image_paths.items():
        if os.path.exists(path):
            unique_files[path] = unique_files.get(path, [])
            unique_files[path].append(name)

    for path, keys in unique_files.items():
        with open(path, 'rb') as f:
            img_data = f.read()
        base64_data = base64.b64encode(img_data).decode('utf-8')
        img_data_uri = f'data:image/png;base64,{base64_data}'

        alt_text = ''
        for k in keys:
            upper_key = k.upper()
            if upper_key in alt_texts:
                alt_text = alt_texts[upper_key]
                break
        if not alt_text:
            alt_text = keys[0] if keys else 'chart'

        for key in keys:
            upper_key = key.upper()
            if not upper_key.isascii():
                continue
            for variant in [upper_key, upper_key.lower(), upper_key.capitalize()]:
                regex_pattern = r'!\[[^\]]*\]\(PLOT:' + re.escape(variant) + r'\)'
                content = re.sub(regex_pattern, f'<img src="{img_data_uri}" alt="{alt_text}">', content)

            compact_regex = r'\[PLOT:' + re.escape(upper_key) + r'\]'
            content = re.sub(compact_regex, f'<img src="{img_data_uri}" alt="{alt_text}">', content)

    html += '<div>' + content + '</div></body></html>'

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)

    return output_path


def create_zip_package(output_dir, zip_path):
    """Create ZIP package with all reports and charts."""
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(output_dir):
            for file in files:
                file_path = os.path.join(root, file)
                if file == 'report.zip':
                    continue
                arcname = os.path.relpath(file_path, output_dir)
                zipf.write(file_path, arcname)

    return zip_path


def main():
    parser = argparse.ArgumentParser(description="Assemble reports from data, charts, and markdown")
    parser.add_argument("--data", "-d", required=True, help="Input CSV data file")
    parser.add_argument("--charts", "-c", required=True, help="Directory containing chart images")
    parser.add_argument("--report", "-r", required=True, help="Markdown report content (or file path)")
    parser.add_argument("--output", "-o", required=True, help="Output directory for reports")
    args = parser.parse_args()

    # Validate inputs
    if not os.path.exists(args.data):
        print(f"Error: Data file not found: {args.data}", file=sys.stderr)
        sys.exit(1)

    if not os.path.isdir(args.charts):
        print(f"Error: Charts directory not found: {args.charts}", file=sys.stderr)
        sys.exit(1)

    # Create output directory
    os.makedirs(args.output, exist_ok=True)

    # Load or use report markdown
    if os.path.exists(args.report):
        with open(args.report, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
    else:
        markdown_text = args.report

    # Load data for potential embedding
    df = pd.read_csv(args.data)

    # Collect chart paths with proper placeholder mapping
    image_paths = {}
    chart_mapping = {
        'DISTRIBUTION': 'score_distribution.png',
        'COMPARISON': 'class_comparison.png',
        'RADAR': 'radar_chart.png',
        'TREND': 'trend_line.png',
        'NORMAL': 'normal_dist.png',
        'BOXPLOT': 'class_boxplot.png',
        'HEATMAP': 'correlation_heatmap.png',
        'DEVIATION': 'deviation_bar.png',
        'TOP_BOTTOM': 'top_bottom_bar.png',
        'SCATTER': 'scatter_regression.png',
        'CDF': 'cdf_curve.png',
        'BOXPLOT_SUBJ': 'subj_boxplot.png',
        '散点图': 'scatter_regression.png',
        '累积分布': 'cdf_curve.png',
        '科目箱线图': 'subj_boxplot.png',
        '分数分布图': 'score_distribution.png',
        '分组对比图': 'class_comparison.png',
        '雷达图': 'radar_chart.png',
        '成绩趋势图': 'trend_line.png',
        '正态分布图': 'normal_dist.png',
        '分组稳定性图': 'class_boxplot.png',
        '科目相关性图': 'correlation_heatmap.png',
        '学生偏离度图': 'deviation_bar.png',
        '尖子生对比图': 'top_bottom_bar.png',
    }

    for chart_key, chart_file in chart_mapping.items():
        chart_path = os.path.join(args.charts, chart_file)
        if os.path.exists(chart_path):
            image_paths[chart_key] = chart_path

    print(f"Found {len(image_paths)} charts")

    # Generate Word report
    word_path = os.path.join(args.output, 'report.docx')
    print(f"Creating Word report: {word_path}")
    create_word_report(markdown_text, image_paths, word_path)
    print(f"Word report saved")

    # Generate HTML report
    html_path = os.path.join(args.output, 'report.html')
    print(f"Creating HTML report: {html_path}")
    create_html_report(markdown_text, image_paths, html_path)
    print(f"HTML report saved")

    # Create ZIP package
    zip_path = os.path.join(args.output, 'report.zip')
    print(f"Creating ZIP package: {zip_path}")
    create_zip_package(args.output, zip_path)
    print(f"ZIP package created")

    print(f"\nAll reports saved to: {args.output}")
    print(f"ZIP package: {zip_path}")


if __name__ == "__main__":
    main()